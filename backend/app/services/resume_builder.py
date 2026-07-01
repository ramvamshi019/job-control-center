"""
services/resume_builder.py
--------------------------
Builds a FULL, ATS-friendly tailored résumé for a single job, on demand.

HONESTY IS THE WHOLE POINT. The builder starts from one of YOUR real base
résumés (resumes/base_*.md) and only:
  - reorders sections/bullets so the JD-relevant experience leads
  - rephrases existing bullets using the posting's terminology ONLY where that
    wording still accurately describes work you actually did
  - drops bullets irrelevant to this role
  - writes a 2-3 line summary assembled from experience already in the résumé
It NEVER invents employers, dates, titles, metrics, skills, or achievements.
Contact info, company names, dates, and education are copied verbatim.

If no Anthropic key is configured (or the call fails), it falls back to the
base résumé unchanged — which is already 100% true. So the worst case is
"honest but un-tailored", never "tailored but fabricated".

The result is written to resumes/generated/<id>_<company>.md (human-readable)
and a stripped .txt (single-column plain text — the format ATS parsers read
most reliably). Returns those paths plus the résumé text.
"""

from __future__ import annotations

import re
from pathlib import Path

import json

from app.config import settings
from app.models.job import Job
from app.services import latex_resume, match_report
from app.services.resume_tailor import pick_base_resume
from app.utils.logging import get_logger

log = get_logger("resume_builder")

# resumes/ and resumes/generated/ live at the project root (this file is
# backend/app/services/resume_builder.py → parents[3] is the repo root).
RESUMES_DIR = Path(__file__).resolve().parents[3] / "resumes"
GENERATED_DIR = RESUMES_DIR / "generated"

# Sources you can apply to WITHOUT creating an account — worth a tailored résumé
# because applying is one click. workday + icims force account creation, so they
# are deliberately excluded (the builder still works on them, just flagged).
LOGIN_FREE_SOURCES = {
    "greenhouse", "lever", "ashby", "bamboohr", "smartrecruiters",
    "recruitee", "workable",
}


def is_login_free(source: str) -> bool:
    return (source or "").lower() in LOGIN_FREE_SOURCES


# Your real résumé, dropped in as resumes/base_master.md, is the source of truth
# for ALL tailoring when present — it overrides the synthetic role-based bases.
MASTER_BASE = "base_master"


def resolve_base(title: str = "") -> str:
    """Use the master (your real) résumé if it exists, else fall back to the
    role-based base picked from the job title."""
    if (RESUMES_DIR / f"{MASTER_BASE}.md").exists():
        return MASTER_BASE
    return pick_base_resume(title)


def _slug(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", (text or "").lower()).strip("-")[:40] or "company"


def _load_base(base: str) -> str:
    """Read a base résumé and strip its internal editor scaffolding (the
    "(Base Resume)" title suffix and the "> Primary resume… never invent"
    blockquote note) so neither the fallback file nor the AI prompt carries
    meta-text into a résumé you'd actually submit."""
    path = RESUMES_DIR / f"{base}.md"
    try:
        raw = path.read_text(encoding="utf-8")
    except Exception as exc:  # noqa: BLE001
        log.warning("could not read base résumé %s: %s", path, exc)
        return ""
    lines = []
    for line in raw.splitlines():
        if line.lstrip().startswith(">"):          # drop the editor note
            continue
        line = re.sub(r"\s*\(Base Resume\)", "", line)  # clean the H1 title
        lines.append(line)
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip() + "\n"


def _md_to_text(md: str) -> str:
    """Strip Markdown to single-column plain text — the most ATS-parser-safe
    form. No tables, no decoration, standard bullets."""
    out = []
    for line in md.splitlines():
        s = line.rstrip()
        s = re.sub(r"^\s*#{1,6}\s*", "", s)          # headings → plain line
        s = re.sub(r"^\s*[-*+]\s+", "- ", s)          # normalize bullets
        s = re.sub(r"^\s*>\s?", "", s)                # blockquotes
        s = s.replace("**", "").replace("__", "")     # bold
        s = re.sub(r"`([^`]*)`", r"\1", s)            # inline code
        s = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", s)  # links → text (url)
        out.append(s)
    return "\n".join(out).strip() + "\n"


def _add_runs(paragraph, text: str) -> None:
    """Render inline **bold** within a docx paragraph; rest as normal runs."""
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if i % 2 == 1:  # odd chunks were inside ** **
            run.bold = True


def _to_docx(resume_md: str, path: Path) -> bool:
    """Write the résumé as a clean, single-column, ATS-friendly .docx — the
    format most application forms parse best. Returns False (and writes nothing)
    if python-docx isn't available, so callers degrade gracefully."""
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:  # noqa: BLE001
        log.warning("python-docx unavailable, skipping .docx: %s", exc)
        return False

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    for raw in resume_md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            level = min(len(h.group(1)), 3)
            doc.add_heading(re.sub(r"\*\*", "", h.group(2)), level=level)
            continue
        b = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if b:
            _add_runs(doc.add_paragraph(style="List Bullet"), b.group(1))
            continue
        _add_runs(doc.add_paragraph(), re.sub(r"^\s*>\s?", "", line))
    doc.save(str(path))
    return True


def _build_prompt(job: Job, base_text: str, missing_keywords: list[str]) -> str:
    missing_line = (
        "The job asks for these keywords your résumé does NOT yet use: "
        f"{', '.join(missing_keywords)}. You may weave one in ONLY if that exact skill "
        "already appears in the base résumé; otherwise leave it out.\n"
        if missing_keywords else
        "No specific missing keywords were detected.\n"
    )
    return (
        "You are an expert résumé editor. Tailor the candidate's BASE RÉSUMÉ to the job "
        "below. This is a strict EDITING task, not a writing task.\n\n"
        "HARD RULES — breaking any of these makes the output unusable:\n"
        "1. NEVER introduce a skill, tool, technology, framework, cloud service, "
        "language, employer, title, date, degree, certification, or metric that is not "
        "already written in the base résumé. If the job wants React/GCP/Flask/etc. and "
        "the base résumé doesn't list it, the résumé does NOT get it. No exceptions. "
        "An honest 80%% beats a fabricated 100%%.\n"
        "2. Copy contact info, employer names, the candidate's OWN PAST JOB TITLES, "
        "employment dates, and education EXACTLY as written. NEVER relabel a held role "
        "to match the target job — a 'Data Engineer' role stays 'Data Engineer'. (The "
        "personal headline/tagline may be reordered among roles the candidate already lists.)\n"
        "3. You MAY: reorder sections and bullets so the most relevant experience leads; "
        "drop bullets irrelevant to this role; rephrase a real bullet using the job's "
        "terminology ONLY when that wording still truthfully describes the SAME work with "
        "the SAME tools the candidate already used; write a 2-3 line professional summary "
        "assembled strictly from experience already present.\n"
        "4. Keep the work-authorization line truthful and unchanged in meaning.\n"
        "5. PRESERVE THE EXACT MARKDOWN STRUCTURE of the base résumé so it renders "
        "cleanly: '# Name', the '**Title:** / **Location:** / **Phone:** · **Email:** · "
        "**LinkedIn:**' header block, '## Section' headers, '### Role — Company · Dates' "
        "then '*Location*' then '-' bullets. No tables, columns, images, or emojis.\n\n"
        f"{missing_line}"
        "Return ONLY a JSON object (no markdown fence, no preamble) with exactly two keys:\n"
        '  "resume_markdown": the full tailored résumé as a Markdown string,\n'
        '  "added_keywords": an array of base-résumé keywords you moved into a more '
        "prominent position for this role (must already exist in the base). Empty array is fine.\n\n"
        f"=== JOB ===\nTitle: {job.title}\nCompany: {job.company_name}\n"
        f"Location: {job.location}\n\nDescription:\n{(job.description or '')[:4000]}\n\n"
        f"=== BASE RÉSUMÉ (the source of truth — every claim must trace back here) ===\n"
        f"{base_text}"
    )


def _call_claude(prompt: str) -> str:
    import anthropic

    client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
    msg = client.messages.create(
        model=settings.resume_model,
        max_tokens=2500,
        messages=[{"role": "user", "content": prompt}],
    )
    return "".join(b.text for b in msg.content if b.type == "text").strip()


def _ai_build(job: Job, base_text: str, missing_keywords: list[str]) -> tuple[str, list[str]]:
    """Tailor the base résumé to the job via Claude, then ENFORCE honesty with a
    deterministic guard: any hard skill in the output that is not in the base résumé
    is a fabrication. If the model invents skills, re-prompt once naming them; if it
    still cheats, fall back to the untouched (100%% true) base. Returns
    (resume_markdown, added_keywords); ('', []) signals 'use the base'."""
    allowed = match_report.extract_skills(base_text)
    prompt = _build_prompt(job, base_text, missing_keywords)
    try:
        md, _added = _parse_ai_json(_call_claude(prompt))
        if not md:
            return "", []
        violations = sorted(match_report.extract_skills(md) - allowed)
        if not violations:
            return md, []  # clean: no skill exists that isn't in the base

        log.warning("résumé for '%s' invented skills not in base: %s — re-prompting",
                    job.title, violations)
        fix = prompt + (
            "\n\n=== CRITICAL CORRECTION ===\nYour previous output INVENTED these skills "
            f"that do NOT appear in the base résumé: {', '.join(violations)}. These are "
            "fabrications and are forbidden. Remove every mention of each (skills lists, "
            "bullets, summary, project tech lines, headers) and output the corrected JSON."
        )
        md2, _ = _parse_ai_json(_call_claude(fix))
        if md2 and not (match_report.extract_skills(md2) - allowed):
            return md2, []
        log.warning("résumé for '%s' still fabricating after retry; using honest base", job.title)
        return "", []
    except Exception as exc:  # noqa: BLE001
        log.warning("AI résumé build failed for '%s', falling back to base: %s",
                    job.title, exc)
        return "", []


def _parse_ai_json(raw: str) -> tuple[str, list[str]]:
    """Pull resume_markdown + added_keywords out of the model's reply, tolerating
    a stray ```json fence or surrounding prose. If the résumé text can't be found,
    return ('', []) so the caller falls back to the honest base."""
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text).strip()
    try:
        obj = json.loads(text)
        md = (obj.get("resume_markdown") or "").strip()
        added = [str(k).strip() for k in (obj.get("added_keywords") or []) if str(k).strip()]
        if md:
            return md, added
    except Exception:  # noqa: BLE001 — fall through to a salvage attempt
        m = re.search(r'"resume_markdown"\s*:\s*"(.*?)"\s*,\s*"added_keywords"', text, re.S)
        if m:
            try:
                md = json.loads('"' + m.group(1) + '"')
            except Exception:  # noqa: BLE001
                md = m.group(1).encode().decode("unicode_escape")
            if md.strip():
                return md.strip(), []
    # Last resort: the model ignored JSON and returned the résumé directly.
    if text.lstrip().startswith("#"):
        return text, []
    return "", []


def build(job: Job, missing_keywords: list[str] | None = None) -> tuple[str, bool, list[str]]:
    """Return (resume_markdown, used_ai, added_keywords). Falls back to the
    untouched base résumé (already true) when AI is off or fails."""
    base = resolve_base(job.title)
    base_text = _load_base(base)
    if settings.resume_ai_enabled and base_text:
        ai, added = _ai_build(job, base_text, missing_keywords or [])
        if ai:
            return ai, True, added
    return base_text, False, []


def save(job: Job, resume_md: str) -> dict:
    """Write the résumé to resumes/generated/ as .md (readable) and .txt (the
    plain single-column form ATS parsers handle best). Returns paths + text."""
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    stem = f"{job.id}_{_slug(job.company_name)}"
    md_path = GENERATED_DIR / f"{stem}.md"
    txt_path = GENERATED_DIR / f"{stem}.txt"
    docx_path = GENERATED_DIR / f"{stem}.docx"
    tex_path = GENERATED_DIR / f"{stem}.tex"
    pdf_path = GENERATED_DIR / f"{stem}.pdf"
    resume_txt = _md_to_text(resume_md)
    md_path.write_text(resume_md, encoding="utf-8")
    txt_path.write_text(resume_txt, encoding="utf-8")
    has_docx = _to_docx(resume_md, docx_path)
    # Jake's Resume LaTeX → PDF (the polished, ATS-friendly upload artifact).
    tex = latex_resume.render_tex(resume_md)
    tex_path.write_text(tex, encoding="utf-8")
    has_pdf = latex_resume.compile_pdf(tex, pdf_path)
    return {
        "md_path": str(md_path),
        "txt_path": str(txt_path),
        "docx_path": str(docx_path) if has_docx else None,
        "tex_path": str(tex_path),
        "pdf_path": str(pdf_path) if has_pdf else None,
        "resume": resume_md,
        "resume_txt": resume_txt,
    }


def get_profile() -> dict:
    """Parse the standard contact/identity fields from the base data résumé so
    the dashboard can offer a one-click copyable autofill kit. Falls back to
    whatever it can find; never raises."""
    text = _load_base(resolve_base("data engineer"))
    def grab(pattern: str) -> str:
        m = re.search(pattern, text, re.I)
        return (m.group(1).strip() if m else "")
    name = ""
    m = re.search(r"^#\s*(.+)$", text, re.M)
    if m:
        name = re.split(r"[—\-|]", m.group(1))[0].strip()
    return {
        "name": name,
        "email": grab(r"Email:\**\s*([^\s·|]+@[^\s·|]+)"),
        "phone": grab(r"Phone:\**\s*([0-9()+\-.\s]{7,})"),
        "linkedin": grab(r"(linkedin\.com/[^\s·|)]+)"),
        "location": grab(r"Location:\**\s*([^·\n|]+)"),
        "work_authorization": settings.my_work_auth,
        "target_roles": ", ".join(settings.target_roles_list),
        "top_skills": ", ".join(settings.skills_list[:12]),
    }


def load_saved(job_id: int) -> dict | None:
    """Return the previously-built résumé for a job (the durable reference you
    keep 'in case you get a call'), or None if none was built. Reads the files
    in resumes/generated/ keyed by job id."""
    mds = sorted(GENERATED_DIR.glob(f"{job_id}_*.md"))
    if not mds:
        return None
    md_path = mds[0]
    txt_path = md_path.with_suffix(".txt")
    try:
        resume_md = md_path.read_text(encoding="utf-8")
    except Exception:  # noqa: BLE001
        return None
    resume_txt = txt_path.read_text(encoding="utf-8") if txt_path.exists() else _md_to_text(resume_md)
    docx_path = md_path.with_suffix(".docx")
    pdf_path = md_path.with_suffix(".pdf")
    tex_path = md_path.with_suffix(".tex")
    return {
        "md_path": str(md_path),
        "txt_path": str(txt_path),
        "docx_path": str(docx_path) if docx_path.exists() else None,
        "tex_path": str(tex_path) if tex_path.exists() else None,
        "pdf_path": str(pdf_path) if pdf_path.exists() else None,
        "resume": resume_md,
        "resume_txt": resume_txt,
    }


def match_report_for(job: Job) -> dict:
    """Jobscan-style skill-coverage report of this job's JD against your master
    résumé — no AI, no generation. Powers the dashboard's 'before' number and the
    have/missing keyword lists."""
    base_text = _load_base(resolve_base(job.title))
    return match_report.analyze(job.title or "", job.description or "", base_text)


def build_and_save(job: Job) -> dict:
    # 1. Score the JD against the base résumé (the honest 'before').
    base_text = _load_base(resolve_base(job.title))
    before = match_report.analyze(job.title or "", job.description or "", base_text)

    # 2. Tailor — telling the model exactly which JD keywords are missing so it can
    #    surface any the candidate's real experience supports (never fabricate).
    resume_md, used_ai, added = build(job, match_report.missing_for_prompt(before))
    if not resume_md:
        raise RuntimeError(
            "No base résumé found in resumes/ — expected base_data_engineer.md etc."
        )

    # 3. Re-score the tailored résumé (the 'after') so the lift is verifiable.
    after = match_report.analyze(job.title or "", job.description or "", resume_md)

    out = save(job, resume_md)
    out["used_ai"] = used_ai
    out["added_keywords"] = added
    out["match_before"] = before["score"]
    out["match_after"] = after["score"]
    out["have"] = after["matched"]
    out["missing"] = after["missing"]
    out["login_free"] = is_login_free(job.source)
    out["base"] = resolve_base(job.title)
    return out
